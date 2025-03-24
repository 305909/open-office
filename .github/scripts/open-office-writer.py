#!/usr/bin/env python3

"""
DOCX Assignment Evaluator for the OpenOffice Writer Course at "I.I.S. G. Cena di Ivrea".

Author: Francesco Giuseppe Gillio
Date: 25.02.2025
"""

import os
import io
import sys
import csv
import ast
import json
import difflib
import zipfile
import subprocess
from PIL import Image
from lxml import etree
from docx import Document


def _get_class_register(registry_path: str) -> dict:
  
    try:
        with open(registry_path, "r") as file:
            registry = json.load(file)
            return registry
    except Exception as e:
        print(f"Unable to load registry {registry_path}: {e}")
        sys.exit(1)


def _get_docx(file_path: str) -> str:
  
    if not file_path.lower().endswith(".odt"):
        return file_path
    docx_file_path = file_path.rsplit('.', 1)[0] + '.docx'
    try:
        subprocess.run(
            [
                'libreoffice',
                '--headless',
                '--convert-to',
                'docx',
                file_path,
                '--outdir',
                os.path.dirname(file_path)
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        return docx_file_path
    except Exception as e:
        print(f"Unable to convert ODT file {file_path}: {e}")
        return file_path


class DocumentAnalyzer:

    def __init__(self, file_path: str):

        self.docx_path = file_path
        try:
            self.doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Unable to load DOCX file '{file_path}': {e}")

    def _get_paragraph_alignment(self, paragraph) -> str:
    
        alignment_map = {0: "left", 1: "center", 2: "right", 3: "justified"}
        try:
            alignment = paragraph.alignment
            return alignment_map.get(alignment, "unknown")
        except AttributeError:
            return "unknown"

    def _get_paragraphs_info(self) -> tuple[list[dict], int]:

        paragraphs_data = []
        empty_lines_count = 0

        for paragraph in self.doc.paragraphs:
            try:
                text = paragraph.text.strip()
            except Exception:
                text = ""
            if not text:
                empty_lines_count += 1
            try:
                paragraph_info = {
                    "text": text,
                    "length": len(text),
                    "style": paragraph.style.name if paragraph.style else "unknown",
                    "bold": any(getattr(run, "bold", False) for run in paragraph.runs),
                    "italic": any(getattr(run, "italic", False) for run in paragraph.runs),
                    "underline": any(getattr(run, "underline", False) for run in paragraph.runs),
                    "font": [run.font.name for run in paragraph.runs if run.font and run.font.name],
                    "size": [run.font.size.pt for run in paragraph.runs if run.font and run.font.size],
                    "alignment": self._get_paragraph_alignment(paragraph)
                }
            except Exception as e:
                paragraph_info = {"text": text, "error": f"Unable to process paragraph {paragraph}: {e}"}
            paragraphs_data.append(paragraph_info)

        return paragraphs_data, empty_lines_count

    def _get_images_info(self) -> list[dict]:

        images_data = []
        for rel in self.doc.part.rels:
            if "image" in self.doc.part.rels[rel].target_ref:
                try:
                    image_part = self.doc.part.rels[rel].target_part
                    image_stream = io.BytesIO(image_part.blob)
                    with Image.open(image_stream) as img:
                        images_data.append({
                            "format": img.format,
                            "dimensions": img.size
                        })
                except Exception as e:
                    images_data.append({
                        "error": f"Unable to process image in relation {rel}: {e}"
                    })
        return images_data

    def _get_tables_info(self) -> list[dict]:

        tables_data = []
        for table in self.doc.tables:
            try:
                num_rows = len(table.rows)
                num_columns = len(table.columns)
                tables_data.append({
                    "rows": num_rows,
                    "columns": num_columns
                })
            except Exception as e:
                tables_data.append({
                    "error": f"Unable to process table {table}: {e}"
                })
        return tables_data

    def _get_margins(self) -> dict:

        margins_data = {}
        try:
            with zipfile.ZipFile(self.docx_path, "r") as docx_zip:
                with docx_zip.open("word/document.xml") as xml_file:
                    xml_tree = etree.parse(xml_file)
                    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    page_margins = xml_tree.xpath("//w:sectPr/w:pgMar", namespaces=namespaces)
                    if page_margins:
                        margins_data = {attr: page_margins[0].get(attr) for attr in page_margins[0].keys()}
        except Exception as e:
            margins_data = {"error": f"Unable to extract margins: {e}"}
        return margins_data


class DocumentComparer:

    def __init__(self, reference_file: str, test_file: str, student_name: str, config: dict = None):

        self.reference_analyzer = DocumentAnalyzer(reference_file)
        self.test_analyzer = DocumentAnalyzer(test_file)
        self.student_name = student_name
        self.config = config

    def _assign_score(self, correct: int, total: int) -> float:

        return (correct / total) * 100.0 if total != 0 else 100.0

    def _compare_elements(self, reference_elements: list, test_elements: list, element_name: str) -> tuple[list[str], float]:

        differences_list = []
        total_elements = len(reference_elements)

        if total_elements == 0 and len(test_elements) == 0:
            return differences_list, 100.0
        if total_elements == 0 or len(test_elements) == 0:
            return differences_list, 0.0

        sum_score = 0.0
        for index, (ref_elem, test_elem) in enumerate(zip(reference_elements, test_elements)):
            element_score = 0.0
          
            if element_name.lower() == "image":
                tol = self.config.get("tolerances", {}).get("image_dimension_tolerance", 0)
                sub_total = 3
                sub_matches = 0
                if ref_elem.get("format") == test_elem.get("format"):
                    sub_matches += 1
                ref_dim = ref_elem.get("dimensions")
                test_dim = test_elem.get("dimensions")
                if ref_dim and test_dim and len(ref_dim) == 2 and len(test_dim) == 2:
                    if abs(ref_dim[0] - test_dim[0]) <= tol:
                        sub_matches += 1
                    if abs(ref_dim[1] - test_dim[1]) <= tol:
                        sub_matches += 1
                element_score = (sub_matches / sub_total) * 100

                if element_score < 100:
                    diff_entry = [
                        f"- **Image {index + 1} mismatch:**",
                        "  - **Differences:**"
                    ]
                    if ref_elem.get("format") != test_elem.get("format"):
                        diff_entry.append("    - **Format:**")
                        diff_entry.append(f"      - **Reference:** {ref_elem.get('format')}")
                        diff_entry.append(f"      - **Student Submission:** {test_elem.get('format')}")
                    if ref_dim and test_dim:
                        if abs(ref_dim[0] - test_dim[0]) > tol:
                            diff_entry.append("    - **Width:**")
                            diff_entry.append(f"      - **Reference:** {ref_dim[0]}")
                            diff_entry.append(f"      - **Student Submission:** {test_dim[0]}")
                        if abs(ref_dim[1] - test_dim[1]) > tol:
                            diff_entry.append("    - **Height:**")
                            diff_entry.append(f"      - **Reference:** {ref_dim[1]}")
                            diff_entry.append(f"      - **Student Submission:** {test_dim[1]}")
                    differences_list.append("\n".join(diff_entry))

            elif element_name.lower() == "table":
                tol_rows = self.config.get("tolerances", {}).get("table_rows_tolerance", 0)
                tol_cols = self.config.get("tolerances", {}).get("table_columns_tolerance", 0)
                sub_total = 2
                sub_matches = 0
                ref_rows = ref_elem.get("rows")
                test_rows = test_elem.get("rows")
                if ref_rows is not None and test_rows is not None:
                    if abs(ref_rows - test_rows) <= tol_rows:
                        sub_matches += 1
                ref_cols = ref_elem.get("columns")
                test_cols = test_elem.get("columns")
                if ref_cols is not None and test_cols is not None:
                    if abs(ref_cols - test_cols) <= tol_cols:
                        sub_matches += 1
                element_score = (sub_matches / sub_total) * 100

                if element_score < 100:
                    diff_entry = [
                        f"- **Table {index + 1} mismatch:**",
                        "  - **Differences:**"
                    ]
                    if ref_rows is not None and test_rows is not None:
                        if abs(ref_rows - test_rows) > tol_rows:
                            diff_entry.append("    - **Rows:**")
                            diff_entry.append(f"      - **Reference:** {ref_rows}")
                            diff_entry.append(f"      - **Student Submission:** {test_rows}")
                    if ref_cols is not None and test_cols is not None:
                        if abs(ref_cols - test_cols) > tol_cols:
                            diff_entry.append("    - **Columns:**")
                            diff_entry.append(f"      - **Reference:** {ref_cols}")
                            diff_entry.append(f"      - **Student Submission:** {test_cols}")
                    differences_list.append("\n".join(diff_entry))

            elif element_name.lower() == "margins":
                tol_margin_cm = self.config.get("tolerances", {}).get("margin_tolerance", 0)
                tol_margin = tol_margin_cm * 1000
                sub_keys = list(ref_elem.keys())
                sub_total = len(sub_keys)
                sub_matches = 0
                for key in sub_keys:
                    try:
                        ref_val = int(ref_elem.get(key))
                        test_val = int(test_elem.get(key))
                        if abs(ref_val - test_val) <= tol_margin:
                            sub_matches += 1
                    except Exception:
                        pass
                element_score = (sub_matches / sub_total) * 100 if sub_total > 0 else 100

                if element_score < 100:
                    diff_entry = [
                        f"- **Margins mismatch:**",
                        "  - **Differences:**"
                    ]
                    for key in sub_keys:
                        try:
                            ref_val = int(ref_elem.get(key))
                            test_val = int(test_elem.get(key))
                            if abs(ref_val - test_val) > tol_margin:
                                diff_entry.append(f"    - **{key.capitalize()}**:")
                                diff_entry.append(f"      - **Reference:** {ref_val}")
                                diff_entry.append(f"      - **Student Submission:** {test_val}")
                        except Exception as e:
                            diff_entry.append(f"    - **{key.capitalize()}**: error comparing values: {e}")
                    differences_list.append("\n".join(diff_entry))

            else:
                match = (ref_elem == test_elem)
                element_score = 100 if match else 0
                if not match:
                    diff_entry = [
                        f"- **{element_name.capitalize()} {index + 1} mismatch:**",
                        "  - **Differences:**"
                    ]
                    if isinstance(ref_elem, dict) and isinstance(test_elem, dict):
                        for key in ref_elem.keys():
                            if key in test_elem and ref_elem[key] != test_elem[key]:
                                diff_entry.append(f"    - **{key.capitalize()}**:")
                                diff_entry.append(f"      - **Reference:** {ref_elem[key]}")
                                diff_entry.append(f"      - **Student Submission:** {test_elem[key]}")
                    else:
                        diff_entry.append(f"    - **Reference:** {ref_elem}")
                        diff_entry.append(f"    - **Student Submission:** {test_elem}")
                    differences_list.append("\n".join(diff_entry))
            sum_score += element_score

        overall_score = sum_score / total_elements
        return differences_list, overall_score

    def _compare_paragraphs(self, ref_paragraphs: list[dict], test_paragraphs: list[dict]) -> tuple[list[str], float]:

        differences = []
        total_score = 0
        count_ref = len(ref_paragraphs)

        if count_ref == 0 and len(test_paragraphs) == 0:
            return differences, 100.0
        if count_ref == 0 or len(test_paragraphs) == 0:
            return differences, 0.0

        min_count = min(count_ref, len(test_paragraphs))
        for i in range(min_count):
            ref_p = ref_paragraphs[i]
            test_p = test_paragraphs[i]

            text_similarity = difflib.SequenceMatcher(None, ref_p.get("text", ""), test_p.get("text", "")).ratio()
            formatting_attributes = ["style", "bold", "italic", "underline", "alignment", "font", "size"]
            formatting_matches = sum(
                1 for attr in formatting_attributes if ref_p.get(attr) == test_p.get(attr)
            )
            formatting_similarity = formatting_matches / len(formatting_attributes)

            paragraph_score = 0.5 * text_similarity + 0.5 * formatting_similarity
            total_score += paragraph_score * 100
            threshold = self.config.get("tolerances", {}).get("paragraph_similarity_threshold", 0.8)

            if paragraph_score < threshold:
                diff_entry = [
                    f"- **Paragraph {i + 1} mismatch:**",
                    f"  - **Text Similarity:** {text_similarity * 100:.1f}%",
                    f"  - **Format Similarity:** {formatting_similarity * 100:.1f}%",
                    "  - **Differences:**"
                ]
                for attr in ["text", "length", "style", "bold", "italic", "underline", "alignment", "font", "size"]:
                    ref_val = ref_p.get(attr)
                    test_val = test_p.get(attr)
                    if ref_val != test_val:
                        diff_entry.append(f"    - **{attr.capitalize()}**:")
                        diff_entry.append(f"      - **Reference:** {ref_val}")
                        diff_entry.append(f"      - **Student Submission:** {test_val}")
                differences.append("\n".join(diff_entry))

        extra = count_ref - min_count
        if extra > 0:
            differences.append(f"{extra} additional paragraph(s) in reference with no match in the test.")

        average_score = total_score / count_ref

        _, ref_empty = self.reference_analyzer._get_paragraphs_info()
        _, test_empty = self.test_analyzer._get_paragraphs_info()
        tolerance_empty = self.config.get("tolerances", {}).get("empty_lines", 1)
        if abs(ref_empty - test_empty) <= tolerance_empty:
            bonus = self.config.get("tolerances", {}).get("paragraph_bonus", 10)
            average_score = min(average_score + bonus, 100)

        return differences, average_score

    def _write_markdown_report(self, report_lines: list[str]) -> str:

        markdown_report = f"# Evaluation Report for {self.student_name}\n\n"
        
        for line in report_lines:
            if line.startswith("Paragraphs:"):
                markdown_report += f"## Paragraphs\n**Score:** {line.split(':', 1)[1].strip()}\n\n"
            elif line.startswith("Images:"):
                markdown_report += f"## Images\n**Score:** {line.split(':', 1)[1].strip()}\n\n"
            elif line.startswith("Tables:"):
                markdown_report += f"## Tables\n**Score:** {line.split(':', 1)[1].strip()}\n\n"
            elif line.startswith("Margins:"):
                markdown_report += f"## Margins\n**Score:** {line.split(':', 1)[1].strip()}\n\n"
            elif line.startswith("Final Score:"):
                markdown_report += f"## Final Score\n**{line}**\n\n"
            elif line.startswith("- **Paragraph"):
                markdown_report += f"{line}\n"
            elif line.startswith("- **Image") or line.startswith("- **Table") or line.startswith("- **Margin"):
                markdown_report += f"{line}\n"
            elif "additional paragraph" in line:
                markdown_report += f"- **{line.strip()}**\n"
            elif line.startswith("  - **Differences:") or line.startswith("    - **"):
                markdown_report += f"{line}\n"
            else:
                markdown_report += f"- {line}\n"

        return markdown_report

    def _compare_documents(self) -> tuple[str, float]:

        report_lines = []

        ref_paragraphs, _ = self.reference_analyzer._get_paragraphs_info()
        test_paragraphs, _ = self.test_analyzer._get_paragraphs_info()
        para_differences, paragraph_score = self._compare_paragraphs(ref_paragraphs, test_paragraphs)
        report_lines.append(f"Paragraphs: {paragraph_score:.1f}% match")
        report_lines.extend(para_differences)

        image_differences, image_score = self._compare_elements(
            self.reference_analyzer._get_images_info(),
            self.test_analyzer._get_images_info(),
            "image"
        )
        report_lines.append(f"Images: {image_score:.1f}% match")
        report_lines.extend(image_differences)

        table_differences, table_score = self._compare_elements(
            self.reference_analyzer._get_tables_info(),
            self.test_analyzer._get_tables_info(),
            "table"
        )
        report_lines.append(f"Tables: {table_score:.1f}% match")
        report_lines.extend(table_differences)

        margin_differences, margin_score = self._compare_elements(
            [self.reference_analyzer._get_margins()],
            [self.test_analyzer._get_margins()],
            "margins"
        )
        report_lines.append(f"Margins: {margin_score:.1f}% match")
        report_lines.extend(margin_differences)

        weights = self.config.get("weights", {})
        w_paragraphs = weights.get("paragraphs", 0.25)
        w_images = weights.get("images", 0.25)
        w_tables = weights.get("tables", 0.25)
        w_margins = weights.get("margins", 0.25)
        final_score = (
            w_paragraphs * paragraph_score +
            w_images * image_score +
            w_tables * table_score +
            w_margins * margin_score
        )
        final_score = min(final_score, 100)
        report_lines.append(f"\nFinal Score: {final_score:.1f}%")

        return self._write_markdown_report(report_lines), final_score


class DocumentEvaluator:

    def __init__(self, assignment_identifier: str, config: dict = None, registry_path: str = None):

        self.assignment_id = assignment_identifier
        self.registry = _get_class_register(registry_path)
        self.assignments_dir = "assignments"
        self.solutions_dir = "solutions"
        self.evaluations_dir = os.path.join("evaluations", self.assignment_id)
        self.assignment_folder = os.path.join(self.assignments_dir, *self.assignment_id.split("/"))

        solution_odt = os.path.join(self.solutions_dir, *self.assignment_id.split("/"), "solution.odt")
        solution_docx = os.path.join(self.solutions_dir, *self.assignment_id.split("/"), "solution.docx")
        if os.path.exists(solution_odt):
            self.solution_file = _get_docx(solution_odt)
        elif os.path.exists(solution_docx):
            self.solution_file = solution_docx
        else:
            print(f"Error: Reference solution file in '{os.path.join(self.solutions_dir, self.assignment_id)}' is not available.")
            sys.exit(1)

        self.report_file = os.path.join(self.evaluations_dir, f"REPORT.csv")
        self.config = config

    def _verify_resources(self) -> bool:

        if not os.path.exists(self.assignment_folder):
            print(f"Error: Assignment folder '{self.assignment_folder}' is not available.")
            return False

        if not os.path.exists(self.solution_file):
            print(f"Error: Reference solution file '{self.solution_file}' is not available.")
            return False

        os.makedirs(self.evaluations_dir, exist_ok=True)
        return True

    def _run_evaluation(self) -> None:

        if not self._verify_resources():
            return

        evaluation_results = []

        for student_id, student_name in self.registry.items():
            student_name_parts = student_name.upper().split()

            matched_file = None
            for file in os.listdir(self.assignment_folder):
                if file.endswith(".docx") or file.endswith(".odt"):
                    file_name = os.path.splitext(os.path.basename(file))[0].upper()
                    file_name_parts = file_name.split()
                    num_words = len(file_name_parts)
                    x_student_name = " ".join(student_name_parts[:num_words])
                    if file_name == x_student_name:
                        matched_file = file
                        print(f"Submission for {student_name}: {file}")
                        break
                      
            if matched_file:
                student_file_path = os.path.join(self.assignment_folder, matched_file)
                student_file_path = _get_docx(student_file_path)
                try:
                    comparer = DocumentComparer(self.solution_file, student_file_path, student_name, config=self.config)
                    report, final_score = comparer._compare_documents()
                except Exception as e:
                    print(f"Execution error for {matched_file}: {e}")
                    continue
            else:
                final_score = 0.0
                report = f"# Report for {student_name}\n\nNo submission, score: {final_score}%\n"
                print(f"No submission for {student_name}.")

            evaluation_results.append({
                "Student": student_name,
                "Score (%)": round(final_score, 2)
            })
          
            individual_report_path = os.path.join(self.evaluations_dir, f"{student_name}.md")
            try:
                with open(individual_report_path, "w") as report_file:
                    report_file.write(report)
            except Exception as e:
                print(f"Unable to write .MD report for {file}: {e}")
            
            print(f"Evaluation for {student_name}: {final_score:.1f}%")
        try:
            with open(self.report_file, "w", newline="") as csvfile:
                fieldnames = ["Student", "Score (%)"]
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                sorted_results = sorted(evaluation_results, key=lambda x: x["Student"])
                for result in sorted_results:
                    writer.writerow(result)
            print(f"Overall Evaluation Report available at: {self.report_file}")
        except Exception as e:
            print(f"Unable to write .CSV report: {e}")


def main():

    if len(sys.argv) < 3:
        print("Error: Please provide the assignment identifier and the registry file path as input parameters.")
        return

    assignment_identifier = sys.argv[1]
    registry_path = sys.argv[2]
    config = None
  
    if len(sys.argv) >= 4:
        config_file = sys.argv[3]
        try:
            with open(config_file, "r") as cf:
                config = json.load(cf)
                print(f"Runtime configuration: {config}")
        except Exception as e:
            print(f"Unable to read configuration file '{config_file}': {e}")
            return

    evaluator = DocumentEvaluator(assignment_identifier, config=config, registry_path=registry_path)
    evaluator._run_evaluation()


if __name__ == "__main__":
    main()
