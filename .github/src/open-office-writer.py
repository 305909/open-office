#!/usr/bin/env python3
"""
DOCX Assignment Evaluator for the OpenOffice Writer Course at "I.I.S. G. Cena di Ivrea".

This module evaluates DOCX assignments by comparing student submissions with a reference solution.
It supports external configuration for weights and tolerances, and automatically converts ODT files
to DOCX using LibreOffice in headless mode.

External parameters:
  - A JSON configuration file with weights and tolerance values.
  - A reference subfolder to extract files for evaluation.

Author: Francesco Giuseppe Gillio
Date: 25.02.2025
"""

import os
import io
import sys
import csv
import ast
import json
import difflib
import zipfile
import subprocess
from PIL import Image
from lxml import etree
from docx import Document


def convert_odt_to_docx(file_path: str) -> str:
    """
    Convert an ODT file to DOCX format using LibreOffice in headless mode.
    
    If the provided file is not in ODT format, the original file path is returned.
    
    Args:
        file_path (str): The path to the file that may be in ODT format.
        
    Returns:
        str: The path to the converted DOCX file or the original file path on error.
    """
    if not file_path.lower().endswith(".odt"):
        return file_path
    # Create the DOCX file path by replacing the extension
    docx_file_path = file_path.rsplit('.', 1)[0] + '.docx'
    try:
        # Run LibreOffice in headless mode to perform the conversion
        subprocess.run(
            [
                'libreoffice',
                '--headless',
                '--convert-to',
                'docx',
                file_path,
                '--outdir',
                os.path.dirname(file_path)
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        return docx_file_path
    except Exception as e:
        print(f"Error converting file {file_path}: {e}")
        return file_path


class DocumentAnalyzer:
    """
    Extracts and parses various components of a DOCX document.
    
    The analyzer extracts:
      - Paragraph details: text content, length, style, font attributes (bold, italic, underline),
        font names, font sizes, and paragraph alignment.
      - Image details: file format and dimensions.
      - Table details: the number of rows and columns.
      - Margin configuration: extracted from the underlying XML structure of the document.
    
    Dependencies:
      - python-docx: for reading and interacting with DOCX files.
      - lxml: for parsing XML content within the DOCX archive.
      - Pillow: for processing and analyzing image files.
    """

    def __init__(self, file_path: str):
        """
        Initialize the DocumentAnalyzer with the provided DOCX file.
        
        Args:
            file_path (str): The path to the DOCX file for analysis.
        
        Raises:
            ValueError: If the DOCX file cannot be loaded.
        """
        self.docx_path = file_path
        try:
            self.doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Unable to load DOCX file '{file_path}': {e}")

    def get_paragraph_alignment(self, paragraph) -> str:
        """
        Determine the alignment of a given paragraph.
        
        This method maps the numerical alignment attribute from a DOCX paragraph
        to a human-readable string. If the alignment cannot be determined, it returns "unknown".
        
        Args:
            paragraph: A paragraph object from the DOCX document.
        
        Returns:
            str: One of "left", "center", "right", "justified", or "unknown".
        """
        alignment_map = {0: "left", 1: "center", 2: "right", 3: "justified"}
        try:
            alignment = paragraph.alignment
            return alignment_map.get(alignment, "unknown")
        except AttributeError:
            return "unknown"

    def get_paragraphs_info(self) -> tuple[list[dict], int]:
        """
        Extract detailed information from each paragraph in the document.
        
        For every paragraph, this method retrieves:
          - The stripped text.
          - The length of the text.
          - The style name.
          - Flags indicating whether any part of the paragraph is bold, italic, or underlined.
          - A list of font names and sizes for each text run.
          - The paragraph alignment.
        
        Additionally, it counts the number of empty paragraphs.
        
        Returns:
            tuple: A tuple containing:
                - A list of dictionaries with detailed paragraph information.
                - An integer count of empty paragraphs.
        """
        paragraphs_data = []
        empty_lines_count = 0

        for paragraph in self.doc.paragraphs:
            try:
                text = paragraph.text.strip()
            except Exception:
                text = ""
            if not text:
                empty_lines_count += 1
            try:
                paragraph_info = {
                    "text": text,
                    "length": len(text),
                    "style": paragraph.style.name if paragraph.style else "unknown",
                    "bold": any(getattr(run, "bold", False) for run in paragraph.runs),
                    "italic": any(getattr(run, "italic", False) for run in paragraph.runs),
                    "underline": any(getattr(run, "underline", False) for run in paragraph.runs),
                    "font": [run.font.name for run in paragraph.runs if run.font and run.font.name],
                    "size": [run.font.size.pt for run in paragraph.runs if run.font and run.font.size],
                    "alignment": self.get_paragraph_alignment(paragraph)
                }
            except Exception as e:
                paragraph_info = {"text": text, "error": f"Error processing paragraph: {e}"}
            paragraphs_data.append(paragraph_info)

        return paragraphs_data, empty_lines_count

    def get_images_info(self) -> list[dict]:
        """
        Retrieve details about images embedded in the DOCX document.
        
        This method iterates over the document's relationships to locate images.
        For each image found, it uses Pillow to determine the image format and dimensions.
        
        Returns:
            list: A list of dictionaries, each containing:
                - "format": The image format (e.g., JPEG, PNG).
                - "dimensions": A tuple representing (width, height).
        """
        images_data = []
        for rel in self.doc.part.rels:
            if "image" in self.doc.part.rels[rel].target_ref:
                try:
                    image_part = self.doc.part.rels[rel].target_part
                    image_stream = io.BytesIO(image_part.blob)
                    with Image.open(image_stream) as img:
                        images_data.append({
                            "format": img.format,
                            "dimensions": img.size
                        })
                except Exception as e:
                    images_data.append({
                        "error": f"Error processing image in relation {rel}: {e}"
                    })
        return images_data

    def get_tables_info(self) -> list[dict]:
        """
        Extract information about tables present in the document.
        
        For each table, this method calculates the number of rows and columns.
        
        Returns:
            list: A list of dictionaries with keys "rows" and "columns" representing table dimensions.
        """
        tables_data = []
        for table in self.doc.tables:
            try:
                num_rows = len(table.rows)
                num_columns = len(table.columns)
                tables_data.append({
                    "rows": num_rows,
                    "columns": num_columns
                })
            except Exception as e:
                tables_data.append({
                    "error": f"Error processing table: {e}"
                })
        return tables_data

    def get_margins(self) -> dict:
        """
        Extract margin configurations from the DOCX document using its XML structure.
        
        The DOCX file is opened as a ZIP archive and the "word/document.xml" file is parsed.
        The method locates the "pgMar" element that defines the page margins.
        
        Returns:
            dict: A dictionary containing margin attributes (such as top, bottom, left, right),
                  or an empty dictionary if margins cannot be extracted.
        """
        margins_data = {}
        try:
            with zipfile.ZipFile(self.docx_path, "r") as docx_zip:
                with docx_zip.open("word/document.xml") as xml_file:
                    xml_tree = etree.parse(xml_file)
                    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    page_margins = xml_tree.xpath("//w:sectPr/w:pgMar", namespaces=namespaces)
                    if page_margins:
                        margins_data = {attr: page_margins[0].get(attr) for attr in page_margins[0].keys()}
        except Exception as e:
            margins_data = {"error": f"Error extracting margins: {e}"}
        return margins_data


class DocumentComparer:
    """
    Compares two DOCX documents based on their structure and formatting.
    
    This class leverages DocumentAnalyzer to extract features from both a reference document
    and a test document (e.g., a student's submission). It then compares features such as
    paragraphs, images, tables, and margins and computes a matching percentage for each section.
    A detailed Markdown report summarizing the differences is generated.
    """

    def __init__(self, reference_file: str, test_file: str, config: dict = None):
        """
        Initialize the DocumentComparer with a reference and a test document.
        
        Also loads configuration for weights and tolerance thresholds.
        
        Args:
            reference_file (str): The path to the reference DOCX file.
            test_file (str): The path to the test DOCX file (student submission).
            config (dict, optional): Configuration dictionary with weights and tolerances.
        """
        self.reference_analyzer = DocumentAnalyzer(reference_file)
        self.test_analyzer = DocumentAnalyzer(test_file)
        # Default configuration for section weights and tolerances
        default_config = {
            "weights": {
                "paragraphs": 0.35,
                "images": 0.25,
                "tables": 0.15,
                "margins": 0.25
            },
            "tolerances": {
                "empty_lines": 0,
                "paragraph_similarity_threshold": 1.00,
                "paragraph_bonus": 0, 
                "image_dimension_tolerance": 0, 
                "table_rows_tolerance": 0, 
                "table_columns_tolerance": 0, 
                "margin_tolerance": 0
            }
        }
        self.config = config if config is not None else default_config

    def assign_score(self, correct: int, total: int) -> float:
        """
        Compute the score for a section as the ratio of correct matches to total items.
        
        If there are no elements to compare, a score of 100% is returned.
        
        Args:
            correct (int): Number of matching elements.
            total (int): Total number of elements compared.
        
        Returns:
            float: The calculated percentage score (0 to 100).
        """
        if total == 0:
            return 100.0
        return (correct / total) * 100.0

    def compare_elements(self, reference_elements: list, test_elements: list,
                         element_name: str) -> tuple[list[str], float]:
        """
        Compare two lists of elements (e.g., images, tables, margins) between documents.
        
        For certain element types, tolerances are applied (e.g., image dimensions, table rows/columns,
        and margin values) to compute a proportional match percentage. For other elements, an exact
        comparison is performed.
        
        Args:
            reference_elements (list): Elements extracted from the reference document.
            test_elements (list): Elements extracted from the test document.
            element_name (str): Descriptive name of the element type.
        
        Returns:
            tuple: A tuple containing:
                - A list of strings describing the differences.
                - A match percentage score for the element type.
        """
        differences_list = []
        total_elements = len(reference_elements)

        # Handle cases where both lists are empty
        if total_elements == 0 and len(test_elements) == 0:
            return differences_list, 100.0
        if total_elements == 0 or len(test_elements) == 0:
            return differences_list, 0.0

        sum_score = 0.0
        for index, (ref_elem, test_elem) in enumerate(zip(reference_elements, test_elements)):
            element_score = 0.0
            if element_name.lower() == "image":
                # Tolerance for image dimension differences
                tol = self.config.get("tolerances", {}).get("image_dimension_tolerance", 0)
                sub_total = 3  # Compare format, width, and height
                sub_matches = 0
                # Exact comparison for image format
                if ref_elem.get("format") == test_elem.get("format"):
                    sub_matches += 1
                # Compare image dimensions (width and height) within the tolerance threshold
                ref_dim = ref_elem.get("dimensions")
                test_dim = test_elem.get("dimensions")
                if ref_dim and test_dim and len(ref_dim) == 2 and len(test_dim) == 2:
                    if abs(ref_dim[0] - test_dim[0]) <= tol:
                        sub_matches += 1
                    if abs(ref_dim[1] - test_dim[1]) <= tol:
                        sub_matches += 1
                element_score = (sub_matches / sub_total) * 100

                # Record detailed differences if image comparison is not perfect
                if element_score < 100:
                    diff_entry = [
                        f"- **Image {index + 1} mismatch:**",
                        "  - **Differences:**"
                    ]
                    if ref_elem.get("format") != test_elem.get("format"):
                        diff_entry.append("    - **Format:**")
                        diff_entry.append(f"      - **Reference:** {ref_elem.get('format')}")
                        diff_entry.append(f"      - **Student Submission:** {test_elem.get('format')}")
                    if ref_dim and test_dim:
                        if abs(ref_dim[0] - test_dim[0]) > tol:
                            diff_entry.append("    - **Width:**")
                            diff_entry.append(f"      - **Reference:** {ref_dim[0]}")
                            diff_entry.append(f"      - **Student Submission:** {test_dim[0]}")
                        if abs(ref_dim[1] - test_dim[1]) > tol:
                            diff_entry.append("    - **Height:**")
                            diff_entry.append(f"      - **Reference:** {ref_dim[1]}")
                            diff_entry.append(f"      - **Student Submission:** {test_dim[1]}")
                    differences_list.append("\n".join(diff_entry))

            elif element_name.lower() == "table":
                # Tolerances for table rows and columns differences
                tol_rows = self.config.get("tolerances", {}).get("table_rows_tolerance", 0)
                tol_cols = self.config.get("tolerances", {}).get("table_columns_tolerance", 0)
                sub_total = 2  # Compare rows and columns
                sub_matches = 0
                ref_rows = ref_elem.get("rows")
                test_rows = test_elem.get("rows")
                if ref_rows is not None and test_rows is not None:
                    if abs(ref_rows - test_rows) <= tol_rows:
                        sub_matches += 1
                ref_cols = ref_elem.get("columns")
                test_cols = test_elem.get("columns")
                if ref_cols is not None and test_cols is not None:
                    if abs(ref_cols - test_cols) <= tol_cols:
                        sub_matches += 1
                element_score = (sub_matches / sub_total) * 100

                # Record table differences if any
                if element_score < 100:
                    diff_entry = [
                        f"- **Table {index + 1} mismatch:**",
                        "  - **Differences:**"
                    ]
                    if ref_rows is not None and test_rows is not None:
                        if abs(ref_rows - test_rows) > tol_rows:
                            diff_entry.append("    - **Rows:**")
                            diff_entry.append(f"      - **Reference:** {ref_rows}")
                            diff_entry.append(f"      - **Student Submission:** {test_rows}")
                    if ref_cols is not None and test_cols is not None:
                        if abs(ref_cols - test_cols) > tol_cols:
                            diff_entry.append("    - **Columns:**")
                            diff_entry.append(f"      - **Reference:** {ref_cols}")
                            diff_entry.append(f"      - **Student Submission:** {test_cols}")
                    differences_list.append("\n".join(diff_entry))

            elif element_name.lower() == "margins":
                # Tolerance for margin differences
                tol_margin = self.config.get("tolerances", {}).get("margin_tolerance", 0)
                sub_keys = list(ref_elem.keys())
                sub_total = len(sub_keys)
                sub_matches = 0
                for key in sub_keys:
                    try:
                        ref_val = int(ref_elem.get(key))
                        test_val = int(test_elem.get(key))
                        if abs(ref_val - test_val) <= tol_margin:
                            sub_matches += 1
                    except Exception:
                        pass
                element_score = (sub_matches / sub_total) * 100 if sub_total > 0 else 100

                # Record margin differences if any
                if element_score < 100:
                    diff_entry = [
                        f"- **Margins mismatch:**",
                        "  - **Differences:**"
                    ]
                    for key in sub_keys:
                        try:
                            ref_val = int(ref_elem.get(key))
                            test_val = int(test_elem.get(key))
                            if abs(ref_val - test_val) > tol_margin:
                                diff_entry.append(f"    - **{key.capitalize()}**:")
                                diff_entry.append(f"      - **Reference:** {ref_val}")
                                diff_entry.append(f"      - **Student Submission:** {test_val}")
                        except Exception as e:
                            diff_entry.append(f"    - **{key.capitalize()}**: error comparing values: {e}")
                    differences_list.append("\n".join(diff_entry))

            else:
                # For other element types, perform an exact equality check.
                match = (ref_elem == test_elem)
                element_score = 100 if match else 0
                if not match:
                    diff_entry = [
                        f"- **{element_name.capitalize()} {index + 1} mismatch:**",
                        "  - **Differences:**"
                    ]
                    if isinstance(ref_elem, dict) and isinstance(test_elem, dict):
                        for key in ref_elem.keys():
                            if key in test_elem and ref_elem[key] != test_elem[key]:
                                diff_entry.append(f"    - **{key.capitalize()}**:")
                                diff_entry.append(f"      - **Reference:** {ref_elem[key]}")
                                diff_entry.append(f"      - **Student Submission:** {test_elem[key]}")
                    else:
                        diff_entry.append(f"    - **Reference:** {ref_elem}")
                        diff_entry.append(f"    - **Student Submission:** {test_elem}")
                    differences_list.append("\n".join(diff_entry))

            sum_score += element_score

        overall_score = sum_score / total_elements
        return differences_list, overall_score

    def compare_paragraphs(self, ref_paragraphs: list[dict], test_paragraphs: list[dict]) -> tuple[list[str], float]:
        """
        Compare paragraphs from two documents based on text and formatting similarity.
        
        For each pair of paragraphs:
          - Compute text similarity using difflib.SequenceMatcher.
          - Compute formatting similarity based on style, bold, italic, underline, alignment, font, and size.
          - Combine the two with equal weighting (50% text and 50% formatting).
        
        If the overall similarity for a paragraph is below a configured threshold,
        the discrepancy is noted. Additionally, extra paragraphs in the reference
        document are penalized, while a bonus may be applied if the empty paragraph counts
        are within tolerance.
        
        Args:
            ref_paragraphs (list[dict]): List of paragraph details from the reference document.
            test_paragraphs (list[dict]): List of paragraph details from the test document.
        
        Returns:
            tuple: A tuple containing:
                - A list of strings detailing the differences.
                - An overall paragraph match percentage.
        """
        differences = []
        total_score = 0
        count_ref = len(ref_paragraphs)

        if count_ref == 0 and len(test_paragraphs) == 0:
            return differences, 100.0
        if count_ref == 0 or len(test_paragraphs) == 0:
            return differences, 0.0

        min_count = min(count_ref, len(test_paragraphs))
        for i in range(min_count):
            ref_p = ref_paragraphs[i]
            test_p = test_paragraphs[i]
            
            # Compute text similarity using difflib's ratio method
            text_similarity = difflib.SequenceMatcher(None, ref_p.get("text", ""), test_p.get("text", "")).ratio()
            # Define formatting attributes to compare
            formatting_attributes = ["style", "bold", "italic", "underline", "alignment", "font", "size"]
            # Count how many formatting attributes match exactly
            formatting_matches = sum(
                1 for attr in formatting_attributes if ref_p.get(attr) == test_p.get(attr)
            )
            formatting_similarity = formatting_matches / len(formatting_attributes)
            
            # Overall paragraph score: average of text and formatting similarities
            paragraph_score = 0.5 * text_similarity + 0.5 * formatting_similarity
            total_score += paragraph_score * 100  # Convert to percentage
            threshold = self.config.get("tolerances", {}).get("paragraph_similarity_threshold", 0.8)
            
            if paragraph_score < threshold:
                diff_entry = [
                    f"- **Paragraph {i + 1} mismatch:**",
                    f"  - **Text Similarity:** {text_similarity * 100:.1f}%",
                    f"  - **Format Similarity:** {formatting_similarity * 100:.1f}%",
                    "  - **Differences:**"
                ]
                for attr in ["text", "length", "style", "bold", "italic", "underline", "alignment"]:
                    ref_val = ref_p.get(attr)
                    test_val = test_p.get(attr)
                    if ref_val != test_val:
                        diff_entry.append(f"    - **{attr.capitalize()}**:")
                        diff_entry.append(f"      - **Reference:** {ref_val}")
                        diff_entry.append(f"      - **Student Submission:** {test_val}")
                differences.append("\n".join(diff_entry))

        # Note any extra paragraphs present in the reference document
        extra = count_ref - min_count
        if extra > 0:
            differences.append(f"{extra} additional paragraph(s) in reference with no match in the test.")

        average_score = total_score / count_ref

        # Apply a bonus if the difference in the number of empty paragraphs is within tolerance.
        _, ref_empty = self.reference_analyzer.get_paragraphs_info()
        _, test_empty = self.test_analyzer.get_paragraphs_info()
        tolerance_empty = self.config.get("tolerances", {}).get("empty_lines", 1)
        if abs(ref_empty - test_empty) <= tolerance_empty:
            bonus = self.config.get("tolerances", {}).get("paragraph_bonus", 10)
            average_score = min(average_score + bonus, 100)
            
        return differences, average_score

    def generate_markdown_report(self, report_lines: list[str]) -> str:
        """
        Create a Markdown-formatted report from a list of report lines.
        
        The report includes clear section headers and detailed sub-lists for differences.
        
        Args:
            report_lines (list[str]): List of strings representing report lines.
            
        Returns:
            str: A Markdown-formatted report.
        """
        markdown_report = "# Evaluation Report\n\n"
        
        for line in report_lines:
            # Define section headers based on line prefixes
            if line.startswith("Paragraphs:"):
                score = line.split(":", 1)[1].strip()
                markdown_report += "## Paragraphs\n"
                markdown_report += f"**Score:** {score}\n\n"
            elif line.startswith("Images:"):
                score = line.split(":", 1)[1].strip()
                markdown_report += "## Images\n"
                markdown_report += f"**Score:** {score}\n\n"
            elif line.startswith("Tables:"):
                score = line.split(":", 1)[1].strip()
                markdown_report += "## Tables\n"
                markdown_report += f"**Score:** {score}\n\n"
            elif line.startswith("Margins:"):
                score = line.split(":", 1)[1].strip()
                markdown_report += "## Margins\n"
                markdown_report += f"**Score:** {score}\n\n"
            elif line.startswith("Final Score:"):
                markdown_report += f"## Final Score\n**{line}**\n\n"
            # Append details regarding mismatches and differences
            elif line.startswith("- **Paragraph"):
                markdown_report += f"{line}\n"
            elif (line.startswith("- **Image") or
                  line.startswith("- **Table") or
                  line.startswith("- **Margin")):
                markdown_report += f"{line}\n"
            elif "additional paragraph" in line:
                markdown_report += f"- **{line.strip()}**\n"
            elif line.startswith("  - **Differences:**") or line.startswith("    - **"):
                markdown_report += f"{line}\n"
            else:
                markdown_report += f"- {line}\n"

        return markdown_report

    def compare_documents(self) -> tuple[str, float]:
        """
        Compare the reference and test DOCX documents section by section.
        
        This method evaluates:
          - Paragraphs: by assessing text and formatting similarity.
          - Images, Tables, and Margins: by directly comparing extracted elements.
        
        A final weighted score is computed based on configuration weights, and a Markdown report
        is generated summarizing all differences.
        
        Returns:
            tuple: A tuple containing the Markdown report (str) and the final overall score (float).
        """
        report_lines = []

        # --- Compare Paragraphs ---
        ref_paragraphs, _ = self.reference_analyzer.get_paragraphs_info()
        test_paragraphs, _ = self.test_analyzer.get_paragraphs_info()
        para_differences, paragraph_score = self.compare_paragraphs(ref_paragraphs, test_paragraphs)
        report_lines.append(f"Paragraphs: {paragraph_score:.1f}% match")
        report_lines.extend(para_differences)

        # --- Compare Images ---
        image_differences, image_score = self.compare_elements(
            self.reference_analyzer.get_images_info(),
            self.test_analyzer.get_images_info(),
            "image"
        )
        report_lines.append(f"Images: {image_score:.1f}% match")
        report_lines.extend(image_differences)

        # --- Compare Tables ---
        table_differences, table_score = self.compare_elements(
            self.reference_analyzer.get_tables_info(),
            self.test_analyzer.get_tables_info(),
            "table"
        )
        report_lines.append(f"Tables: {table_score:.1f}% match")
        report_lines.extend(table_differences)

        # --- Compare Margins ---
        margin_differences, margin_score = self.compare_elements(
            [self.reference_analyzer.get_margins()],
            [self.test_analyzer.get_margins()],
            "margins"
        )
        report_lines.append(f"Margins: {margin_score:.1f}% match")
        report_lines.extend(margin_differences)

        # --- Compute Final Score using configuration weights ---
        weights = self.config.get("weights", {})
        w_paragraphs = weights.get("paragraphs", 0.25)
        w_images = weights.get("images", 0.25)
        w_tables = weights.get("tables", 0.25)
        w_margins = weights.get("margins", 0.25)
        final_score = (
            w_paragraphs * paragraph_score +
            w_images * image_score +
            w_tables * table_score +
            w_margins * margin_score
        )
        final_score = min(final_score, 100)
        report_lines.append(f"\nFinal Score: {final_score:.1f}%")

        return self.generate_markdown_report(report_lines), final_score


class DocxAssignmentEvaluator:
    """
    Manages the evaluation of student DOCX submissions for a given assignment.
    
    This evaluator locates the proper subfolder for the assignment (both for submissions and solutions),
    compares each student submission against the reference solution using DocumentComparer, and generates
    individual Markdown reports as well as an overall CSV report.
    """

    def __init__(self, assignment_identifier: str, config: dict = None):
        """
        Initialize the evaluator with the assignment identifier and optional configuration.
        
        Args:
            assignment_identifier (str): A unique identifier for the assignment.
            config (dict, optional): A dictionary containing weights and tolerance values.
        """
        self.assignment_id = assignment_identifier
        self.assignments_dir = "assignments"
        self.solutions_dir = "solutions"
        self.evaluations_dir = os.path.join("evaluations", self.assignment_id)

        self.assignment_folder = os.path.join(self.assignments_dir, *self.assignment_id.split("/"))

        # Determine the reference solution file (convert ODT to DOCX if necessary)
        solution_odt = os.path.join(self.solutions_dir, *self.assignment_id.split("/"), "solution.odt")
        solution_docx = os.path.join(self.solutions_dir, *self.assignment_id.split("/"), "solution.docx")
        if os.path.exists(solution_odt):
            self.solution_file = convert_odt_to_docx(solution_odt)
        elif os.path.exists(solution_docx):
            self.solution_file = solution_docx
        else:
            print(f"Error: Reference solution file in '{os.path.join(self.solutions_dir, self.assignment_id)}' is not available.")
            sys.exit(1)

        self.report_file = os.path.join(self.evaluations_dir, f"Report.csv")
        self.config = config

    def verify_resources(self) -> bool:
        """
        Verify the existence of necessary folders and files for the evaluation process.
        
        Returns:
            bool: True if all required resources exist; otherwise, False.
        """
        if not os.path.exists(self.assignment_folder):
            print(f"Error: Assignment folder '{self.assignment_folder}' is not available.")
            return False

        if not os.path.exists(self.solution_file):
            print(f"Error: Reference solution file '{self.solution_file}' is not available.")
            return False

        os.makedirs(self.evaluations_dir, exist_ok=True)
        return True

    @staticmethod
    def extract_student_name(file_path: str) -> tuple[str, str]:
        """
        Extract the student's first name and surname from the DOCX file name.
        
        Assumes the file name is formatted as "firstname-surname.docx".
        
        Args:
            file_path (str): The path to the student's DOCX file.
        
        Returns:
            tuple: A tuple (first_name, surname).
        """
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        name_parts = base_name.split('-')
        if len(name_parts) >= 2:
            first_name, surname = [part.lower().capitalize() for part in name_parts[:2]]
        else:
            first_name, surname = base_name.lower().capitalize(), ""
        return first_name, surname

    def run_evaluation(self) -> None:
        """
        Execute the evaluation for each student submission in the assignment folder.
        
        For each DOCX (or ODT) file found:
          - Convert ODT files to DOCX if necessary.
          - Compare the submission with the reference solution using DocumentComparer.
          - Generate an individual Markdown report and record the final score.
          - Compile overall results into a CSV report.
        """
        if not self.verify_resources():
            return

        evaluation_results = []

        for file in os.listdir(self.assignment_folder):
            if file.endswith(".docx") or file.endswith(".odt"):
                student_file_path = os.path.join(self.assignment_folder, file)
                # Convert ODT files to DOCX if needed
                student_file_path = convert_odt_to_docx(student_file_path)
                try:
                    comparer = DocumentComparer(self.solution_file, student_file_path, config=self.config)
                    report, final_score = comparer.compare_documents()
                except Exception as e:
                    print(f"Error processing file {file}: {e}")
                    continue

                first_name, surname = self.extract_student_name(student_file_path)
                evaluation_results.append({
                    "Name": first_name,
                    "Surname": surname,
                    "Score (%)": final_score
                })

                individual_report_path = os.path.join(
                    self.evaluations_dir, f"{first_name}-{surname}.md"
                )
                try:
                    with open(individual_report_path, "w") as report_file:
                        report_file.write(report)
                except Exception as e:
                    print(f"Error writing report for file {file}: {e}")

                print(f"Evaluation for {file}: {final_score:.1f}% (report at {individual_report_path})")

        try:
            with open(self.report_file, "w", newline="") as csvfile:
                fieldnames = ["Name", "Surname", "Score (%)"]
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                for result in evaluation_results:
                    writer.writerow(result)
            print(f"Overall Evaluation Report available at: {self.report_file}")
        except Exception as e:
            print(f"Error writing CSV report: {e}")


def main():
    """
    Main function to execute the DOCX assignment evaluation.
    
    The first command-line argument must be the assignment identifier.
    Optionally, a second argument may specify the path to a JSON configuration file with weights and tolerances.
    """
    if len(sys.argv) < 2:
        print("Error: Please provide the assignment identifier as an input parameter.")
        return

    assignment_identifier = sys.argv[1]
    config = None
  
    if len(sys.argv) >= 3:
        config_file = sys.argv[2]
        try:
            with open(config_file, "r") as cf:
                config = json.load(cf)
                print(f"Loaded configuration: {config}")
                if config:
                    weights = config.get("weights", {})
                    tolerances = config.get("tolerances", {})
                    print(f"Using Weights: {weights}")
                    print(f"Using Tolerances: {tolerances}")
        except Exception as e:
            print(f"Error reading configuration file '{config_file}': {e}")
            return

    evaluator = DocxAssignmentEvaluator(assignment_identifier, config=config)
    evaluator.run_evaluation()


if __name__ == "__main__":
    main()
